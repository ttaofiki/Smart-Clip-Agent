"""
Generates Smart_Clip_Agent_Documentation.docx
Run once: python build_docs.py
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

doc = Document()

# ── Page margins ──────────────────────────────────────────────────────────────
for section in doc.sections:
    section.top_margin    = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin   = Cm(2.8)
    section.right_margin  = Cm(2.8)

# ── Style helpers ─────────────────────────────────────────────────────────────
BRAND   = RGBColor(0x1F, 0x53, 0x8D)   # dark blue
ACCENT  = RGBColor(0x27, 0xAE, 0x60)   # green
WARN    = RGBColor(0xE6, 0x7E, 0x22)   # orange
CODE_BG = RGBColor(0xF5, 0xF5, 0xF5)


def h1(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after  = Pt(6)
    run = p.add_run(text)
    run.bold      = True
    run.font.size = Pt(20)
    run.font.color.rgb = BRAND
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '4')
    bottom.set(qn('w:color'), '1F538D')
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p


def h2(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after  = Pt(4)
    run = p.add_run(text)
    run.bold      = True
    run.font.size = Pt(15)
    run.font.color.rgb = BRAND
    return p


def h3(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after  = Pt(2)
    run = p.add_run(text)
    run.bold      = True
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0x34, 0x49, 0x5E)
    return p


def body(text, space_after=4):
    p = doc.add_paragraph(text)
    p.paragraph_format.space_after = Pt(space_after)
    for run in p.runs:
        run.font.size = Pt(11)
    return p


def bullet(text, level=0, bold_prefix=None):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.left_indent  = Inches(0.3 + level * 0.25)
    p.paragraph_format.space_after  = Pt(3)
    if bold_prefix:
        r1 = p.add_run(bold_prefix)
        r1.bold = True
        r1.font.size = Pt(11)
        r2 = p.add_run(text)
        r2.font.size = Pt(11)
    else:
        r = p.add_run(text)
        r.font.size = Pt(11)
    return p


def numbered(text, bold_prefix=None):
    p = doc.add_paragraph(style='List Number')
    p.paragraph_format.left_indent = Inches(0.3)
    p.paragraph_format.space_after = Pt(4)
    if bold_prefix:
        r1 = p.add_run(bold_prefix)
        r1.bold = True
        r1.font.size = Pt(11)
        r2 = p.add_run(text)
        r2.font.size = Pt(11)
    else:
        r = p.add_run(text)
        r.font.size = Pt(11)
    return p


def code_block(*lines):
    for line in lines:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent  = Inches(0.4)
        p.paragraph_format.space_before = Pt(1)
        p.paragraph_format.space_after  = Pt(1)
        run = p.add_run(line if line else " ")
        run.font.name = 'Courier New'
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0x2C, 0x3E, 0x50)
        pPr = p._p.get_or_add_pPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'),   'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'),  'F0F0F0')
        pPr.append(shd)


def note(text, kind="note"):
    colors = {"note": "1F538D", "warning": "E67E22", "tip": "27AE60"}
    labels = {"note": "  NOTE: ", "warning": "  IMPORTANT: ", "tip": "  TIP: "}
    p = doc.add_paragraph()
    p.paragraph_format.left_indent  = Inches(0.3)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)
    color = RGBColor.from_string(colors.get(kind, "1F538D"))
    r1 = p.add_run(labels.get(kind, "NOTE: "))
    r1.bold = True
    r1.font.color.rgb = color
    r1.font.size = Pt(11)
    r2 = p.add_run(text)
    r2.font.size = Pt(11)
    r2.font.color.rgb = RGBColor(0x2C, 0x3E, 0x50)
    return p


def spacer():
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)


def divider():
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after  = Pt(6)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'),   'single')
    bottom.set(qn('w:sz'),    '4')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), 'CCCCCC')
    pBdr.append(bottom)
    pPr.append(pBdr)


def add_table(headers, rows, col_widths=None):
    tbl = doc.add_table(rows=1 + len(rows), cols=len(headers))
    tbl.style = 'Table Grid'
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT

    hdr = tbl.rows[0]
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        cell.text = h
        run = cell.paragraphs[0].runs[0]
        run.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'),   'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'),  '1F538D')
        tcPr.append(shd)

    for r_idx, row_data in enumerate(rows):
        row = tbl.rows[r_idx + 1]
        for c_idx, cell_text in enumerate(row_data):
            cell = row.cells[c_idx]
            cell.text = str(cell_text)
            for run in cell.paragraphs[0].runs:
                run.font.size = Pt(10)
            if r_idx % 2 == 1:
                tc = cell._tc
                tcPr = tc.get_or_add_tcPr()
                shd = OxmlElement('w:shd')
                shd.set(qn('w:val'),   'clear')
                shd.set(qn('w:color'), 'auto')
                shd.set(qn('w:fill'),  'EAF2FB')
                tcPr.append(shd)

    if col_widths:
        for row in tbl.rows:
            for i, width in enumerate(col_widths):
                row.cells[i].width = Inches(width)

    doc.add_paragraph()
    return tbl


# ══════════════════════════════════════════════════════════════════════════════
# TITLE PAGE
# ══════════════════════════════════════════════════════════════════════════════

p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(60)
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Smart Clip Agent")
run.bold = True
run.font.size = Pt(32)
run.font.color.rgb = BRAND

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Full Installation & User Documentation")
run.font.size = Pt(16)
run.font.color.rgb = RGBColor(0x55, 0x6B, 0x80)

spacer()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Powered by Claude AI (Anthropic)  •  faster-whisper  •  ffmpeg  •  yt-dlp")
run.font.size = Pt(11)
run.font.color.rgb = RGBColor(0x7F, 0x8C, 0x8D)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Version 1.1  —  Windows 10 / 11")
run.font.size = Pt(10)
run.font.color.rgb = RGBColor(0xAA, 0xAA, 0xAA)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# TABLE OF CONTENTS
# ══════════════════════════════════════════════════════════════════════════════

h1("Table of Contents")

toc_items = [
    ("1.", "Introduction",                              "3"),
    ("2.", "System Requirements",                       "3"),
    ("3.", "Pre-Installation Checklist",                "4"),
    ("4.", "Installation Guide",                        "5"),
    ("  4.1", "Step 1 — Install Python",               "5"),
    ("  4.2", "Step 2 — Run setup.bat",                "6"),
    ("  4.3", "Step 3 — Set Up Your Anthropic API Key","6"),
    ("  4.4", "Step 4 — Launch the App",               "7"),
    ("5.", "Using the Application (GUI)",               "8"),
    ("  5.1", "Input Source",                          "8"),
    ("  5.2", "Clip Settings",                         "9"),
    ("  5.3", "Running a Job",                         "9"),
    ("  5.4", "Output",                                "9"),
    ("6.", "Using the Command-Line Interface (CLI)",    "10"),
    ("7.", "How the Pipeline Works",                    "11"),
    ("8.", "Project File Structure",                    "13"),
    ("9.", "Configuration & Advanced Options",          "14"),
    ("10.", "Troubleshooting",                         "15"),
    ("11.", "Frequently Asked Questions",              "17"),
]

for num, title, page in toc_items:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(3)
    tab_stops = p.paragraph_format.tab_stops
    tab_stops.add_tab_stop(Inches(5.5))
    r1 = p.add_run(f"{num}  {title}")
    r1.font.size = Pt(11)
    if not num.startswith(" "):
        r1.bold = True
    r2 = p.add_run(f"\t{page}")
    r2.font.size = Pt(11)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 1. INTRODUCTION
# ══════════════════════════════════════════════════════════════════════════════

h1("1. Introduction")

body(
    "Smart Clip Agent is an AI-powered desktop application that automatically identifies and "
    "extracts the most interesting highlight clips from any online video or local video file. "
    "It supports YouTube, TikTok, Instagram, Twitter/X, Facebook, Vimeo, Twitch, and hundreds "
    "of other platforms — as well as local files in any format ffmpeg can read (.mp4, .mkv, "
    ".avi, .mov, .webm, and more). Instead of manually scrubbing through hours of footage to "
    "find the best moments, you simply paste a link (or browse to a file), choose how many clips "
    "you want, and the agent does all the work."
)
spacer()
body("The application uses a multi-step AI pipeline:")
bullet("Downloads the video from the platform (or loads from your disk).")
bullet("Converts all speech to a timestamped transcript using faster-whisper.")
bullet("Analyses the audio for energy peaks and natural silence boundaries.")
bullet("Sends the transcript, audio data, and video metadata to Claude AI (Anthropic), which scores every segment and selects the best highlights.")
bullet("Cuts each highlight into a separate .mp4 file using ffmpeg.")
spacer()
body(
    "The result is a set of ready-to-share .mp4 clips saved in the output/ folder, "
    "each named after the highlight it contains."
)

divider()

# ══════════════════════════════════════════════════════════════════════════════
# 2. SYSTEM REQUIREMENTS
# ══════════════════════════════════════════════════════════════════════════════

h1("2. System Requirements")

body("Before installing, confirm your system meets the following requirements:")
spacer()

add_table(
    ["Component",       "Minimum",          "Recommended"],
    [
        ["Operating System", "Windows 10 (64-bit)", "Windows 11 (64-bit)"],
        ["RAM",              "4 GB",                "8 GB or more"],
        ["Disk Space",       "3 GB free",           "10 GB free (for video files)"],
        ["Internet",         "Required (download + AI API calls)", "Broadband (5 Mbps+)"],
        ["Python",           "3.9",                 "3.11 or 3.12"],
        ["GPU",              "Not required",        "CUDA GPU speeds up transcription"],
    ],
    col_widths=[1.8, 2.2, 2.2]
)

note(
    "The transcription step runs on your CPU by default and can take 1–5 minutes for "
    "a typical 10-minute video. A NVIDIA GPU will speed this up significantly, but is not required.",
    "note"
)

divider()

# ══════════════════════════════════════════════════════════════════════════════
# 3. PRE-INSTALLATION CHECKLIST
# ══════════════════════════════════════════════════════════════════════════════

doc.add_page_break()
h1("3. Pre-Installation Checklist")

body(
    "Complete every item on this checklist before running the installer. "
    "Skipping any step is the most common cause of errors."
)
spacer()

h2("3.1  Things You Need Before Starting")

bullet("A Windows 10 or Windows 11 computer.")
bullet("An internet connection.")
bullet(
    "An Anthropic API key — the AI that scores your highlights requires this. "
    "It is free to sign up and you receive free credits to get started. "
    "See Section 4.3 for instructions on obtaining the key."
)
bullet("Enough free disk space to store the downloaded video (typically 200 MB – 2 GB depending on video length and quality).")
spacer()

h2("3.2  Software Pre-Requisites")

body(
    "The setup.bat script (Section 4.2) installs most dependencies automatically. "
    "However, Python must be installed manually before running it."
)
spacer()

add_table(
    ["Software",    "Who installs it",      "Notes"],
    [
        ["Python 3.9+",      "You — manually (Section 4.1)", "MUST be added to PATH"],
        ["ffmpeg",           "setup.bat (automatic)",        "Video processing tool"],
        ["yt-dlp",           "setup.bat (automatic)",        "Multi-platform video downloader"],
        ["faster-whisper",   "setup.bat (automatic)",        "Speech-to-text library"],
        ["anthropic",        "setup.bat (automatic)",        "Claude AI client library"],
        ["customtkinter",    "setup.bat (automatic)",        "Modern GUI library"],
        ["ffmpeg-python",    "setup.bat (automatic)",        "Python wrapper for ffmpeg"],
        ["numpy",            "setup.bat (automatic)",        "Numerical computing library"],
    ],
    col_widths=[1.8, 2.2, 2.2]
)

note(
    "Python must be installed BEFORE running setup.bat. The installer cannot install Python "
    "for you. Follow the exact steps in Section 4.1.",
    "warning"
)

divider()

# ══════════════════════════════════════════════════════════════════════════════
# 4. INSTALLATION GUIDE
# ══════════════════════════════════════════════════════════════════════════════

doc.add_page_break()
h1("4. Installation Guide")

body(
    "Follow these four steps in order. Each step must complete successfully "
    "before moving to the next."
)

# 4.1
h2("4.1  Step 1 — Install Python")

body("Python is the programming language that powers the AI pipeline. You must install it once on each machine.")
spacer()
numbered("Open your web browser and go to:  https://www.python.org/downloads/")
numbered("Click the large yellow Download Python button (version 3.11 or newer recommended).")
numbered("Run the downloaded installer file (e.g., python-3.11.x-amd64.exe).")
numbered(
    "CRITICAL — On the first screen of the installer, tick the checkbox labelled "
    '"Add Python to PATH" or "Add python.exe to PATH" before clicking Install Now. '
    "If you miss this step, Python will not be found by the application."
)
numbered("Click Install Now and wait for the installation to complete.")
numbered("Click Close when done.")
spacer()

note(
    'To verify the installation worked, open a new PowerShell window and type: '
    'python --version\n'
    'You should see output like: Python 3.11.9\n'
    'If you see an error, Python was not added to PATH — re-run the installer and tick the checkbox.',
    "tip"
)

# 4.2
h2("4.2  Step 2 — Run setup.bat")

body(
    "The project folder contains a file called setup.bat which automatically installs "
    "all remaining dependencies. You only need to run this once."
)
spacer()
numbered("Open File Explorer and navigate to the Smart_Clip_Agent folder.")
numbered("Double-click setup.bat.")
numbered("A black Command Prompt window will open and begin installing packages. This may take 3–10 minutes depending on your internet speed.")
numbered("Read the output as it runs. You will see four numbered stages:")

bullet("[1/4] Checking Python",    level=1)
bullet("[2/4] Installing Python packages  (yt-dlp, faster-whisper, anthropic, etc.)", level=1)
bullet("[3/4] Installing ffmpeg  (video processing tool)", level=1)
bullet("[4/4] Setting up API key file  (creates .env)", level=1)

numbered("When stage [4/4] is reached, Notepad will open automatically with the .env file. Do not close it yet — proceed to Step 3.")
spacer()

note(
    "If setup.bat shows 'winget: command not found' for ffmpeg, it means your Windows version "
    "does not have the winget package manager. In that case, download ffmpeg manually from "
    "https://ffmpeg.org/download.html, extract it, and add the bin\\ folder to your system PATH.",
    "warning"
)

# 4.3
h2("4.3  Step 3 — Set Up Your Anthropic API Key")

body(
    "The application uses Claude AI (made by Anthropic) to intelligently score and select "
    "the best highlight moments. Claude requires an API key to authenticate your requests."
)
spacer()

h3("Getting a free API key:")
numbered("Open your browser and go to:  https://console.anthropic.com")
numbered("Click Sign Up and create a free account.")
numbered("Once logged in, click API Keys in the left sidebar.")
numbered('Click Create Key, give it a name (e.g., "Smart Clip Agent"), and copy the key. It starts with sk-ant-...')
spacer()

h3("Adding the key to the application:")
numbered("The Notepad window that opened in Step 2 shows a file called .env with this content:")
code_block("ANTHROPIC_API_KEY=your_api_key_here")
numbered("Replace your_api_key_here with the key you just copied. Example:")
code_block("ANTHROPIC_API_KEY=sk-ant-api03-AbCdEfGhIjKlMnOpQrStUv...")
numbered("Save the file (Ctrl + S) and close Notepad.")
spacer()

note(
    "Keep your API key private. Do not share the .env file or paste your key into emails, "
    "chat messages, or public code. Treat it like a password.",
    "warning"
)
note(
    "New Anthropic accounts receive free usage credits. Processing a typical 10-minute "
    "video costs approximately $0.01–$0.05 in API credits.",
    "note"
)

# 4.4
h2("4.4  Step 4 — Launch the App")

body("You are now ready to run the application.")
spacer()

h3("Option A — Double-click launcher (recommended for daily use):")
numbered("In the project folder, double-click run.bat.")
numbered("The Smart Clip Agent window will open.")
spacer()

h3("Option B — From a terminal:")
numbered("Open PowerShell or Command Prompt.")
numbered("Navigate to the project folder:")
code_block('cd "C:\\Users\\YourName\\Desktop\\Smart_Clip_Agent"')
numbered("Run:")
code_block("python app.py")
spacer()

note(
    "Always open a new terminal window after running setup.bat. ffmpeg is added to your "
    "PATH during installation, and existing terminal windows do not pick up PATH changes "
    "until they are restarted.",
    "note"
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 5. USING THE GUI
# ══════════════════════════════════════════════════════════════════════════════

h1("5. Using the Application (GUI)")

body(
    "The Graphical User Interface (GUI) is the main way to use Smart Clip Agent. "
    "It is designed to be simple — fill in the form, click Start, and wait for your clips."
)

h2("5.1  Input Source")

body(
    "At the top of the window you will see the INPUT SOURCE panel. "
    "Choose one of two modes:"
)
spacer()

add_table(
    ["Mode",            "When to use",                              "How to use"],
    [
        ["Video URL",
         "You want to download and clip a video from YouTube, TikTok, Instagram, Twitter/X, Vimeo, or any other supported platform.",
         "Select 'YouTube URL', paste the full video URL into the text box. Example: https://www.youtube.com/watch?v=..."],
        ["Local Video File",
         "You already have a video file on your computer (.mp4, .mkv, .avi, .mov, .webm, etc.).",
         "Select 'Local Video File', then click Browse to open a file picker and navigate to your video."],
    ],
    col_widths=[1.5, 2.2, 2.5]
)

note(
    "yt-dlp (the downloader used internally) supports over 1,000 video platforms including "
    "YouTube, TikTok, Instagram Reels, Twitter/X, Facebook, Vimeo, Twitch, Dailymotion, "
    "and many more. If a platform hosts publicly accessible videos, it likely works.",
    "tip"
)

h2("5.2  Clip Settings")

body("Below the input, the CLIP SETTINGS panel controls how the highlights are selected:")
spacer()

add_table(
    ["Setting",              "Default", "Description"],
    [
        ["Number of clips",   "5",   "How many highlight clips to extract. Minimum: 1, Maximum: 20."],
        ["Min duration (s)",  "20",  "Shortest allowed clip in seconds. Clips shorter than this are discarded."],
        ["Max duration (s)",  "120", "Longest allowed clip in seconds. Clips longer than this are discarded. Must be greater than Min duration."],
    ],
    col_widths=[1.8, 0.9, 3.5]
)

note(
    "For short-form social media content (TikTok, Instagram Reels, YouTube Shorts), "
    "try: Min duration = 15s, Max duration = 60s, Number of clips = 5.",
    "tip"
)

h2("5.3  Running a Job")

numbered("Fill in the Input Source and Clip Settings.")
numbered("Click the START CLIPPING button.")
numbered("The button will change to Running... and two progress bars will appear:")

bullet("Overall bar (blue) — shows total pipeline progress from 0% to 100%.", level=1)
bullet("Step bar (green) — shows how far along the current step is (e.g., download at 45%).", level=1)

numbered("The Output Log at the bottom shows real-time messages from each stage.")
numbered("When complete, the output/ folder opens automatically and you will see your .mp4 clips.")
spacer()

note(
    "Do not close the app while it is running. The transcription step (Transcribing speech) "
    "takes the longest — typically 1–5 minutes for a 10-minute video. This is normal.",
    "note"
)

h2("5.4  Output")

body(
    "Finished clips are saved in the output/ folder inside the project directory. "
    "Each clip is named after its highlight title, for example:"
)
code_block(
    "output/",
    "  Best_Opening_Hook_clip01.mp4",
    "  Surprising_Revelation_clip02.mp4",
    "  Emotional_Moment_clip03.mp4",
)
spacer()
body("You can also click the Open Output Folder button at the bottom of the app at any time.")

divider()
doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 6. CLI
# ══════════════════════════════════════════════════════════════════════════════

h1("6. Using the Command-Line Interface (CLI)")

body(
    "Advanced users can also run the pipeline from a terminal without opening the GUI. "
    "This is useful for automation, scripting, or running on a server."
)
spacer()

h2("6.1  Basic Usage")

code_block(
    "cd C:\\Users\\YourName\\Desktop\\Smart_Clip_Agent",
    "",
    "# From a video URL (YouTube, TikTok, Instagram, etc.):",
    'node main.js --url "https://www.youtube.com/watch?v=..."',
    'node main.js --url "https://www.tiktok.com/@user/video/..."',
    "",
    "# From a local video file:",
    'node main.js --file "C:\\Users\\YourName\\Videos\\lecture.mp4"',
    'node main.js --file "C:\\Users\\YourName\\Videos\\clip.mkv"',
)
spacer()

note("Node.js is required for the CLI (main.js). The GUI (app.py) does not need Node.js.", "note")

h2("6.2  CLI Options")

add_table(
    ["Option",           "Default", "Description"],
    [
        ["--url <url>",          "—",    "Video URL from any supported platform to download and process."],
        ["--file <path>",        "—",    "Path to a local video file (.mp4, .mkv, .avi, .mov, .webm, etc.). Use instead of --url."],
        ["--clips <N>",          "5",    "Number of highlight clips to extract (1–20)."],
        ["--min-duration <S>",   "20",   "Minimum clip length in seconds."],
        ["--max-duration <S>",   "120",  "Maximum clip length in seconds."],
        ["--help",               "—",    "Show usage instructions and exit."],
    ],
    col_widths=[1.9, 0.9, 3.4]
)

h2("6.3  CLI Examples")

body("Extract 3 clips from a YouTube video, each between 30 and 90 seconds:")
code_block('node main.js --url "https://www.youtube.com/watch?v=..." --clips 3 --min-duration 30 --max-duration 90')
spacer()
body("Clip a locally downloaded video into 10 highlights:")
code_block('node main.js --file "C:\\Videos\\conference.mp4" --clips 10')
spacer()
body("Extract 5 short clips from a TikTok video (good for Reels / Shorts):")
code_block('node main.js --url "https://www.tiktok.com/@user/video/..." --clips 5 --min-duration 15 --max-duration 60')
spacer()
body("Process a local .mkv file:")
code_block('node main.js --file "C:\\Videos\\lecture.mkv" --clips 5 --min-duration 20 --max-duration 120')

divider()
doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 7. HOW THE PIPELINE WORKS
# ══════════════════════════════════════════════════════════════════════════════

h1("7. How the Pipeline Works")

body(
    "Understanding the pipeline helps you diagnose issues and get better results. "
    "The agent runs five sequential stages, each building on the previous one."
)
spacer()

add_table(
    ["Stage", "Tool Used",       "What Happens",                             "Output File"],
    [
        ["1. Download",    "yt-dlp",          "Downloads the video from any supported platform (YouTube, TikTok, Instagram, etc.) in best available quality, or copies a local file (.mp4, .mkv, .avi, .mov, .webm, etc.) into the working directory. Also extracts video metadata: title, duration, description, and chapter markers.", "temp/video.*\ntemp/metadata.json"],
        ["2. Transcribe",  "faster-whisper",  "Converts all speech in the video into a timestamped text transcript. Uses the 'small' Whisper model (fast, accurate enough for highlight detection). Also detects the spoken language automatically.", "temp/transcript.json"],
        ["3. Audio Analysis", "ffmpeg",        "Extracts the audio track and computes the loudness (RMS energy) for every 1-second window. Identifies the top 20 highest-energy moments and detects silence gaps (natural cut points).", "temp/audio_analysis.json"],
        ["4. AI Scoring",  "Claude AI",       "Claude receives the full transcript, audio peaks, silence boundaries, video metadata, and chapter markers. It analyses the content and selects the best N highlight segments using a scoring framework (Hook, Energy, Content Value, Emotion, Duration). Returns a ranked list with titles and reasons.", "temp/highlights.json"],
        ["5. Clip Cutting", "ffmpeg",         "Reads the highlight timestamps and cuts each segment from the original video using stream copy (no quality loss). Falls back to re-encoding if stream copy fails. Names each clip after its AI-generated title.", "output/*.mp4"],
    ],
    col_widths=[1.2, 1.3, 3.0, 1.7]
)

h2("7.1  The AI Scoring Framework")

body(
    "Claude scores each candidate segment out of 100 points using five criteria. "
    "Understanding these criteria helps you predict which moments will be selected:"
)
spacer()

add_table(
    ["Criterion",          "Max Points", "Description"],
    [
        ["Hook / Opening",     "20", "Does the segment start with a strong statement, compelling question, or attention-grabbing moment?"],
        ["Audio Energy",       "20", "Does the segment overlap with a high-energy audio peak detected in Stage 3?"],
        ["Content Value",      "20", "Does it contain a key insight, revelation, quotable line, practical tip, or climax of the video?"],
        ["Emotional Resonance","20", "Does it convey humour, surprise, conflict, inspiration, or strong emotion?"],
        ["Duration Fit",       "20", "Does the segment naturally fit between your chosen Min and Max duration, with clean cut points?"],
        ["Chapter Bonus",      "+10","Bonus: if the segment starts within 5 seconds of a chapter boundary, +10 is added."],
    ],
    col_widths=[1.8, 1.0, 3.4]
)

h2("7.2  Intermediate Files")

body(
    "All temporary files are stored in the temp/ folder. You can inspect them to "
    "understand what each stage produced:"
)
spacer()

add_table(
    ["File",                    "Format", "Contents"],
    [
        ["temp/job.json",           "JSON",  "Job configuration: input URL/path, number of clips, min/max duration."],
        ["temp/video.*",            "Video", "The downloaded or copied source video used by all subsequent stages. Extension matches the source format (.mp4, .mkv, etc.)."],
        ["temp/metadata.json",      "JSON",  "Video title, duration (seconds), description, uploader, and chapter list."],
        ["temp/transcript.json",    "JSON",  "Array of speech segments: [{text, start, end, words: [{word, start, end}]}]"],
        ["temp/audio_analysis.json","JSON",  "Energy timeline, top-20 peaks, silence boundaries, and total duration."],
        ["temp/highlights.json",    "JSON",  "Ranked highlight list: [{rank, start, end, duration, score, title, reason}]"],
        ["temp/clips_summary.json", "JSON",  "Summary of created clips: file paths, sizes, durations, scores."],
    ],
    col_widths=[2.0, 0.8, 3.4]
)

divider()
doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 8. PROJECT FILE STRUCTURE
# ══════════════════════════════════════════════════════════════════════════════

h1("8. Project File Structure")

body("The complete directory layout of the Smart Clip Agent project:")
spacer()

code_block(
    "Smart_Clip_Agent/",
    "|",
    "+-- app.py               <- Desktop GUI (run this to launch the app)",
    "+-- main.js              <- Command-line entry point (requires Node.js)",
    "+-- run.bat              <- Double-click to launch the GUI",
    "+-- setup.bat            <- Double-click to install all dependencies",
    "+-- requirements.txt     <- Python package list (used by setup.bat)",
    "+-- package.json         <- Node.js project config",
    "+-- .env                 <- Your API key (never share this file)",
    "+-- .env.example         <- Template for the .env file",
    "+-- CLAUDE.md            <- Agent role definitions (Ruflo orchestration)",
    "|",
    "+-- src/",
    "|   +-- tools/",
    "|       +-- download.py         <- Multi-platform downloader / local file ingestor",
    "|       +-- transcribe.py       <- faster-whisper speech-to-text",
    "|       +-- audio_analysis.py   <- ffmpeg loudness + silence analysis",
    "|       +-- score_highlights.py <- Claude AI highlight scorer",
    "|       +-- clip.py             <- ffmpeg clip cutter",
    "|",
    "+-- temp/                <- Intermediate files (auto-created, can be deleted)",
    "|   +-- video.mp4  (or .mkv, .webm, etc.)",
    "|   +-- metadata.json",
    "|   +-- transcript.json",
    "|   +-- audio_analysis.json",
    "|   +-- highlights.json",
    "|",
    "+-- output/              <- Your final .mp4 clips land here",
    "    +-- Great_Opening_clip01.mp4",
    "    +-- Key_Insight_clip02.mp4",
)

divider()
doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 9. CONFIGURATION
# ══════════════════════════════════════════════════════════════════════════════

h1("9. Configuration & Advanced Options")

h2("9.1  The .env File")

body(
    "The .env file in the project root stores sensitive settings. "
    "Open it in any text editor (Notepad, VS Code, etc.) to change values."
)
spacer()
code_block("ANTHROPIC_API_KEY=sk-ant-...")
spacer()

note("Never commit the .env file to Git or share it publicly. It contains your private API key.", "warning")

h2("9.2  Changing the AI Model")

body(
    "By default, the scorer uses claude-sonnet-4-6 which provides the best balance of "
    "speed and quality. To change the model, open src/tools/score_highlights.py and edit the model line:"
)
code_block(
    '# Current (default):',
    'model="claude-sonnet-4-6"',
    '',
    '# Faster / cheaper alternative:',
    'model="claude-haiku-4-5-20251001"',
    '',
    '# Most powerful (slower, higher cost):',
    'model="claude-opus-4-7"',
)

h2("9.3  Changing the Transcription Model")

body(
    "The 'small' model is used by default — it is fast and accurate enough for highlight detection. "
    "To switch to a higher-accuracy model, open src/tools/transcribe.py and change MODEL_SIZE:"
)
code_block(
    "# Default (fast, recommended):",
    'MODEL_SIZE = "small"',
    "",
    "# Higher accuracy, ~4x slower on CPU:",
    'MODEL_SIZE = "medium"',
    "",
    "# Highest accuracy, much slower:",
    'MODEL_SIZE = "large-v3"',
)
spacer()

add_table(
    ["Model",     "Speed",      "Accuracy",   "Best For"],
    [
        ["tiny",     "Very fast",  "Low",        "Quick tests, short clips"],
        ["small",    "Fast",       "Good",       "All videos (default)"],
        ["medium",   "Moderate",   "Very good",  "When higher accuracy is needed"],
        ["large-v3", "Slow",       "Excellent",  "High-accuracy requirement"],
    ],
    col_widths=[1.2, 1.2, 1.2, 2.6]
)

h2("9.4  Keeping or Clearing Temporary Files")

body(
    "The temp/ folder grows with each run. It is safe to delete its contents between jobs. "
    "The output/ folder contains your finished clips — only delete files from here "
    "when you no longer need them."
)
code_block(
    "# To clear temp files from PowerShell:",
    "Remove-Item temp\\* -Recurse -Force",
)

divider()
doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 10. TROUBLESHOOTING
# ══════════════════════════════════════════════════════════════════════════════

h1("10. Troubleshooting")

h2("10.1  'Python is not recognised as an internal or external command'")

body("Python is not on your system PATH.")
bullet("Re-run the Python installer.", bold_prefix="Fix: ")
bullet('On the first screen, tick "Add Python to PATH" before clicking Install Now.', level=1)
bullet("Open a new terminal after re-installing and try: python --version", level=1)
spacer()

h2("10.2  'ModuleNotFoundError: No module named anthropic' (or any other module)")

body("A required Python package is not installed.")
bullet("Open a terminal in the project folder and run:", bold_prefix="Fix: ")
code_block("pip install -r requirements.txt")
bullet("If a specific module is still missing, install it directly:", level=1)
code_block("pip install anthropic")
spacer()

h2("10.3  'ANTHROPIC_API_KEY not set' or API errors")

body("The .env file is missing or the key is not set correctly.")
bullet("Open the .env file in Notepad.", bold_prefix="Fix: ")
bullet("Confirm the line reads: ANTHROPIC_API_KEY=sk-ant-...", level=1)
bullet("Make sure there are no extra spaces around the = sign.", level=1)
bullet("Make sure the key is not still the placeholder: your_api_key_here", level=1)
spacer()

h2("10.4  ffmpeg is not found / 'No audio data extracted'")

body("ffmpeg was not installed or is not on PATH.")
bullet("Open a new terminal window (PATH changes need a fresh terminal).", bold_prefix="Fix: ")
bullet("Type: ffmpeg -version — if it shows version info, ffmpeg is installed correctly.", level=1)
bullet("If not found, install manually:", level=1)
code_block(
    "# Option 1 — winget (Windows 10/11):",
    "winget install Gyan.FFmpeg",
    "",
    "# Option 2 — manual:",
    "# Download from https://ffmpeg.org/download.html",
    "# Extract to C:\\ffmpeg\\",
    "# Add C:\\ffmpeg\\bin\\ to your system PATH",
)
spacer()

h2("10.5  Download fails / 'ERROR: Video unavailable'")

body(
    "The video may be age-restricted, private, region-blocked, members-only, or the URL "
    "is incorrect. Some platforms (e.g. Instagram private accounts, Twitter DMs) require "
    "authentication which is not supported."
)
bullet("Confirm the video plays in your browser before trying to clip it.", bold_prefix="Check: ")
bullet("Ensure the full URL is copied, including the https:// prefix.", bold_prefix="Check: ")
bullet("For Instagram/TikTok, make sure the account and video are public.", bold_prefix="Check: ")
bullet("Update yt-dlp (platform downloaders need frequent updates):", bold_prefix="Fix: ")
code_block("pip install --upgrade yt-dlp")
spacer()

h2("10.6  Transcription is very slow")

body("The transcription model is running on CPU, which is normal but takes time.")
bullet("For a 10-minute video, expect 2–5 minutes on a modern CPU — this is normal.", bold_prefix="Note: ")
bullet("Closing other applications frees CPU for faster processing.", bold_prefix="Tip: ")
bullet('The app uses the "small" model by default, which is the fastest option. If still slow, check that no other heavy programs are running.', bold_prefix="Tip: ")
spacer()

h2("10.7  Output clips are empty or very small files")

body("Stream-copy failed silently. The fallback re-encoding may also have failed.")
bullet("Check the Output Log in the app for [clip] FAILED messages.", bold_prefix="Check: ")
bullet("Ensure the source video in temp/ is not corrupted by playing it in VLC or Windows Media Player.", bold_prefix="Check: ")
bullet("Try re-running with a different video to isolate whether the issue is video-specific.", bold_prefix="Try: ")
spacer()

h2("10.8  'No highlights.json found — cannot cut clips'")

body("The AI scoring step failed before producing output.")
bullet("Confirm ANTHROPIC_API_KEY is set correctly in .env.", bold_prefix="Check: ")
bullet("Check your Anthropic account has remaining API credits at console.anthropic.com.", bold_prefix="Check: ")
bullet("Check your internet connection — the scorer calls an external API.", bold_prefix="Check: ")

divider()
doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 11. FAQ
# ══════════════════════════════════════════════════════════════════════════════

h1("11. Frequently Asked Questions")

h3("Q: Is it free to use?")
body(
    "The application itself is free. However, the AI highlight scoring uses the Anthropic API, "
    "which has a cost based on usage. New accounts receive free credits. Processing a typical "
    "10-minute video costs approximately $0.01–$0.05. You can monitor your usage at "
    "console.anthropic.com."
)
spacer()

h3("Q: Which platforms are supported for URL downloads?")
body(
    "Smart Clip Agent uses yt-dlp internally, which supports over 1,000 video platforms. "
    "This includes YouTube, TikTok, Instagram Reels and posts (public only), Twitter/X, "
    "Facebook (public videos), Vimeo, Twitch clips and VODs, Dailymotion, Reddit videos, "
    "and many more. If a platform hosts publicly accessible videos, it very likely works."
)
spacer()

h3("Q: What video formats does it support for local files?")
body(
    "Any format that ffmpeg can read, including: "
    ".mp4, .mkv, .avi, .mov, .webm, .flv, .m4v, .ts, .wmv, and many others. "
    "Use the Browse button to navigate to your file — all common formats are shown in the file picker."
)
spacer()

h3("Q: Can I use it without an internet connection?")
body(
    "Partially. If you already have a local video file, the Transcription and Audio Analysis "
    "steps work offline (once the Whisper model is downloaded on first use). However, the AI "
    "Scoring step always requires an internet connection to call the Claude API. "
    "URL downloads also require internet access."
)
spacer()

h3("Q: How do I install it on another computer?")
body("Copy the entire project folder to the other computer, then:")
numbered("Install Python from python.org (add to PATH).")
numbered("Double-click setup.bat.")
numbered("Add the ANTHROPIC_API_KEY to the .env file.")
numbered("Double-click run.bat.")
spacer()

h3("Q: Why does the app open the output/ folder automatically when done?")
body(
    "This is a convenience feature so you can immediately see and play your clips. "
    "If you find this disruptive, you can remove the "
    "self.after(800, lambda: os.startfile(...)) line from app.py."
)
spacer()

h3("Q: The clips are not the moments I wanted. How do I improve results?")
body("The AI selects highlights based on the scoring criteria. To influence the selection:")
bullet("Increase the number of clips — more clips means more coverage of the video.")
bullet("Adjust Min/Max duration — tighter constraints guide the AI to more focused segments.")
bullet("Videos with chapters give the AI additional context and generally produce better results.")
bullet("Videos with clear speech produce better transcripts, which directly improves AI scoring.")
spacer()

h3("Q: Can I run multiple videos in a row?")
body(
    "Yes. After each job completes, the output/ folder retains its clips. "
    "The temp/ folder is overwritten on each new run. Simply paste a new URL or browse "
    "to a new file and click Start again."
)
spacer()

h3("Q: Is my video data sent to Anthropic?")
body(
    "Only text is sent to Anthropic — specifically the video transcript, audio peak timestamps, "
    "silence boundaries, and video metadata (title, description, chapters). "
    "The actual video file is never uploaded to any external server. "
    "All video processing (downloading, transcription, audio analysis, clip cutting) "
    "happens locally on your computer."
)

divider()

# ── Footer ────────────────────────────────────────────────────────────────────
spacer()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Smart Clip Agent  •  Documentation v1.1  •  Built with Claude AI by Anthropic")
run.font.size = Pt(9)
run.font.color.rgb = RGBColor(0xAA, 0xAA, 0xAA)

# ── Save ──────────────────────────────────────────────────────────────────────
out = "Smart_Clip_Agent_Documentation.docx"
doc.save(out)
print(f"Saved: {out}")
